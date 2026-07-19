"""
Build a local .docx of the restaurant leads, formatted the same way
upload_to_google_doc.py would write it into a Google Doc. Produced as a
stopgap while the Google Docs upload is blocked on missing credentials.json —
the user can drag this straight into Google Drive and it auto-converts to a
native Google Doc (Drive > New > File upload > right-click > Open with Google Docs).

Usage:
    python build_local_docx.py --in leads/kenya_restaurants.csv --out leads/kenya_restaurants.docx
"""

from __future__ import annotations

import argparse
import csv
from pathlib import Path

from docx import Document


def build_rows_by_area(rows: list[dict]) -> dict[str, list[dict]]:
    by_area: dict[str, list[dict]] = {}
    for row in rows:
        by_area.setdefault(row.get("query_area", "Other"), []).append(row)
    return by_area


def main():
    parser = argparse.ArgumentParser(description="Build a local .docx of restaurant leads.")
    parser.add_argument("--in", dest="in_path", type=Path, required=True)
    parser.add_argument("--out", dest="out_path", type=Path, required=True)
    parser.add_argument("--title", default="Kenya Restaurant Leads")
    args = parser.parse_args()

    with args.in_path.open(newline="", encoding="utf-8") as f:
        rows = list(csv.DictReader(f))

    by_area = build_rows_by_area(rows)

    doc = Document()
    doc.add_heading(f"{args.title} — {len(rows)} restaurants", level=1)

    for area in sorted(by_area):
        doc.add_heading(f"{area} ({len(by_area[area])})", level=2)
        for r in by_area[area]:
            phone = r.get("phone") or ""
            phone = "" if phone == "N/A" else phone
            bits = [r.get("name", "")]
            if r.get("rating"):
                bits.append(f"★{r['rating']}")
            if phone:
                bits.append(phone)
            if r.get("address_snippet"):
                bits.append(r["address_snippet"])
            doc.add_paragraph(" — ".join(b for b in bits if b), style="List Bullet")

    args.out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out_path)
    print(f"Saved {args.out_path} ({len(rows)} restaurants, {len(by_area)} areas)")


if __name__ == "__main__":
    main()
